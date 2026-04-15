import os

from docx import Document
from openai import OpenAI
from pydantic import BaseModel, Field

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# 1. Define structured output schema
class DocumentUpdate(BaseModel):
    section_header: str = Field(
        description="Exact heading text from the original document"
    )
    action: str = Field(description="replace | append | insert_after")
    new_content: str = Field(description="The exact text to insert or replace with")


class UpdatePlan(BaseModel):
    updates: list[DocumentUpdate] = Field(
        description="Only include sections that need changes"
    )


# 2. Extract document into heading-based sections
def extract_sections(doc_path: str) -> tuple[Document, dict]:
    doc = Document(doc_path)
    sections = {}
    current_header = "Front Matter"
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            current_header = para.text.strip()
            sections[current_header] = {"text": [], "para_ref": para}
        if para.text.strip():
            sections.setdefault(current_header, {"text": [], "para_ref": None})[
                "text"
            ].append(para.text.strip())
    return doc, sections


# 3. Build prompt & call LLM with JSON schema
def get_update_plan(sections: dict, new_instructions: str) -> UpdatePlan:
    context = "\n\n".join(
        [
            f"[{header}] {' '.join(sec['text'])}"
            for header, sec in sections.items()
            if sec["text"]
        ]
    )

    prompt = f"""You are an expert document editor. Given the original document sections and new instructions/text, output a JSON plan of targeted changes.

RULES:
- ONLY include sections that actually need modification.
- Use EXACT heading names from the document.
- Actions: 'replace' (overwrite), 'append' (add to end), 'insert_after' (add as new paragraph)
- Preserve original tone, terminology, and paragraph breaks.
- Output ONLY valid JSON. No markdown, no explanations.

ORIGINAL DOCUMENT:
{context}

NEW TEXT / INSTRUCTIONS:
{new_instructions}
"""

    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={
            "type": "json_schema",
            "json_schema": UpdatePlan.model_json_schema(),
        },
        temperature=0.2,
    )

    return UpdatePlan.model_validate_json(completion.choices[0].message.content)


# 4. Apply changes programmatically
def apply_updates(doc: Document, sections: dict, plan: UpdatePlan) -> Document:
    # Build lookup for fast matching
    header_to_para = {
        h: sec["para_ref"] for h, sec in sections.items() if sec["para_ref"]
    }

    for update in plan.updates:
        header = update.section_header
        if header not in header_to_para:
            print(f"⚠️ Skipping: Heading '{header}' not found in document")
            continue

        ref_para = header_to_para[header]

        if update.action == "replace":
            ref_para.clear()
            ref_para.add_run(update.new_content)

        elif update.action == "append":
            # Append to existing text in the same paragraph
            ref_para.add_run("\n\n" + update.new_content)

        elif update.action == "insert_after":
            # Insert new paragraph after the heading
            new_para = ref_para.insert_paragraph_after(update.new_content)
            new_para.style = (
                ref_para.style.next_paragraph_style
                if ref_para.style.next_paragraph_style
                else "Normal"
            )

    return doc


# 📥 Main Pipeline
def smart_patch_document(input_path: str, output_path: str, new_instructions: str):
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Document not found: {input_path}")

    print("📖 Extracting document structure...")
    doc, sections = extract_sections(input_path)

    print("🤖 Generating update plan...")
    plan = get_update_plan(sections, new_instructions)
    print(f"📋 LLM plans {len(plan.updates)} changes:")
    for u in plan.updates:
        print(f"  • {u.section_header} → {u.action}")

    print("✍️ Applying changes...")
    apply_updates(doc, sections, plan)
    doc.save(output_path)
    print(f"✅ Saved updated document: {output_path}")


# 🔧 Example Usage
if __name__ == "__main__":
    NEW_INSTRUCTIONS = """
    1. In the "Executive Summary", replace the outdated Q3 revenue projection with: 
       'Q4 projections show 18% YoY growth, driven by AI automation adoption.'
    2. Append to "Risk Factors": 'Supply chain delays in APAC region may impact delivery timelines by 2-3 weeks.'
    3. Insert after "Compliance Updates': 'All data handling now complies with EU AI Act Article 14.'
    """

    smart_patch_document("original.docx", "patched.docx", NEW_INSTRUCTIONS)
